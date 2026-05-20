#!/usr/bin/env python3
"""
Markdown to Styled Word Document Converter

Converts structured markdown files to professionally styled Word documents
with branding colors, fonts, and layout. Does not modify content.
"""

import json
import os
import re
import sys
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================================
# LAYER 1: CONFIG & CONSTANTS
# ============================================================================

def load_branding_config(config_path=None):
    """Load branding config from JSON file. Falls back to hardcoded if absent."""
    if config_path is None:
        script_dir = Path(__file__).parent
        config_path = script_dir.parent / "config" / "branding.json"

    if isinstance(config_path, str):
        config_path = Path(config_path)

    if config_path.exists():
        try:
            with open(config_path, 'r') as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            pass

    # Fallback to hardcoded config
    return {
        "colors": {
            "primary_blue": "#0B1EEA",
            "violet": "#3A0CA3",
            "dark_purple": "#240046",
            "white": "#FFFFFF",
            "black": "#000000",
            "light_grey": "#D0D0D0",
            "inline_code_bg": "#E8E8E8",
            "code_block_bg": "#F2F2F2",
            "blockquote_border": "#A0A0A0",
            "blockquote_text": "#646464"
        },
        "callouts": {
            "NOTE": {"border": "#0B1EEA", "fill": "#EEF2FF"},
            "TIP": {"border": "#2D9E4F", "fill": "#EAFAF1"},
            "WARNING": {"border": "#F5A623", "fill": "#FFFBEA"},
            "IMPORTANT": {"border": "#D0342C", "fill": "#FEF2F2"}
        },
        "fonts": {"body": "Calibri", "code": "Courier New", "fallback": "Arial"},
        "sizes": {
            "h1_pt": 28, "h2_pt": 18, "h3_pt": 11, "body_pt": 10,
            "code_pt": 9, "caption_pt": 9, "table_header_pt": 10, "table_body_pt": 10
        },
        "headings": {
            "h1_color": "#000000", "h2_color": "#0B1EEA", "h3_color": "#000000",
            "h1_bold": True, "h2_bold": True, "h3_bold": True
        },
        "table": {
            "header_bg": "#0B1EEA", "header_fg": "#FFFFFF",
            "border_color": "#D0D0D0", "center_table": True
        },
        "page": {"width_inches": 8.5, "height_inches": 11.0, "margin_inches": 1.0}
    }


# Load config at import time
CONFIG = load_branding_config()
COLORS = CONFIG["colors"]
CALLOUT_COLORS = CONFIG["callouts"]
FONTS = CONFIG["fonts"]
SIZES = CONFIG["sizes"]
HEADINGS = CONFIG["headings"]
TABLE_CFG = CONFIG["table"]
PAGE_CFG = CONFIG["page"]

# RGB color objects
PRIMARY_BLUE = RGBColor(11, 30, 234)      # #0B1EEA
VIOLET = RGBColor(58, 12, 163)            # #3A0CA3
DARK_PURPLE = RGBColor(36, 0, 70)         # #240046
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)
LIGHT_GREY = RGBColor(208, 208, 208)      # #D0D0D0


def hex_to_rgb(hex_color):
    """Convert hex color string to RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def hex_to_rgbcolor(hex_color):
    """Convert hex color string to RGBColor."""
    r, g, b = hex_to_rgb(hex_color)
    return RGBColor(r, g, b)


# ============================================================================
# LAYER 2: OOXML PRIMITIVE HELPERS
# ============================================================================

def set_cell_background(cell, fill_color):
    """Set cell background color via OOXML w:shd element."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color.lstrip('#'))
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_run_shading(run, fill_hex):
    """Set character-level background color (for inline code)."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), fill_hex.lstrip('#'))
    run._element.get_or_add_rPr().append(shading_elm)


def set_paragraph_shading(paragraph, fill_hex):
    """Set paragraph-level background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_hex.lstrip('#'))
    paragraph._element.get_or_add_pPr().append(shading_elm)


def set_paragraph_left_border(paragraph, color_hex, sz=24, space=4):
    """Add left border to paragraph (for callouts and blockquotes)."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(sz))
    left.set(qn('w:space'), str(space))
    left.set(qn('w:color'), color_hex.lstrip('#'))
    pBdr.append(left)
    pPr.append(pBdr)


def set_paragraph_indent(paragraph, left_twips=360, right_twips=0):
    """Set paragraph indentation (in twips: 1440 twips = 1 inch)."""
    pPr = paragraph._element.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    if left_twips:
        ind.set(qn('w:left'), str(left_twips))
    if right_twips:
        ind.set(qn('w:right'), str(right_twips))
    pPr.append(ind)


def add_horizontal_rule(doc):
    """Add a horizontal rule (grey divider line)."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), COLORS['light_grey'].lstrip('#'))
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_hyperlink(paragraph, text, url, color=PRIMARY_BLUE):
    """Add clickable hyperlink to paragraph."""
    part = paragraph._element.getparent()
    doc_part = part.getroottree().getroot()

    # Create the hyperlink element
    r = OxmlElement('w:hyperlink')
    r.set(qn('w:tooltip'), text)

    # Create run inside hyperlink
    run_elm = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Underline and color
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    color_elm = OxmlElement('w:color')
    color_elm.set(qn('w:val'), '%02x%02x%02x' % (color[0], color[1], color[2]))
    rPr.append(color_elm)

    run_elm.append(rPr)

    # Add text
    t = OxmlElement('w:t')
    t.text = text
    run_elm.append(t)

    r.append(run_elm)
    paragraph._element.append(r)


def set_table_borders(table):
    """Set table borders: horizontal only, no vertical lines."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')

    # Color for borders
    border_color = TABLE_CFG['border_color'].lstrip('#')

    # Create border elements: top, left, bottom, right, insideH, insideV
    borders_to_create = {
        'w:top': True,
        'w:left': False,
        'w:bottom': True,
        'w:right': False,
        'w:insideH': True,
        'w:insideV': False
    }

    for border_name, show in borders_to_create.items():
        border_elm = OxmlElement(border_name)
        border_elm.set(qn('w:val'), 'single' if show else 'none')
        border_elm.set(qn('w:sz'), '4')
        if show:
            border_elm.set(qn('w:color'), border_color)
        tblBorders.append(border_elm)

    tblPr.append(tblBorders)


# ============================================================================
# LAYER 3: INLINE MARKDOWN PARSER
# ============================================================================

# Compiled regex for inline markdown matching
INLINE_TOKEN_RE = re.compile(
    r'(?P<image>!\[(?P<alt>[^\]]*)\]\((?P<img_url>[^)]+)\))'
    r'|(?P<link>\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)]+)\))'
    r'|(?P<bold>\*\*(?P<bold_text>(?:(?!\*\*).)+?)\*\*)'
    r'|(?P<italic>\*(?P<italic_text>(?:(?!\*).)+?)\*)'
    r'|(?P<code>`(?P<code_text>[^`]+?)`)'
)


def parse_inline(paragraph, text, base_font_size=10):
    """Parse inline markdown (bold, italic, code, links, images) and populate paragraph."""
    if not text or not text.strip():
        return

    last_end = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        # Add plain text before this match
        if match.start() > last_end:
            gap = text[last_end:match.start()]
            r = paragraph.add_run(gap)
            r.font.size = Pt(base_font_size)
            r.font.name = FONTS['body']
            r.font.color.rgb = BLACK

        # Handle match
        if match.group('bold'):
            r = paragraph.add_run(match.group('bold_text'))
            r.font.bold = True
            r.font.size = Pt(base_font_size)
            r.font.name = FONTS['body']
            r.font.color.rgb = BLACK

        elif match.group('italic'):
            r = paragraph.add_run(match.group('italic_text'))
            r.font.italic = True
            r.font.size = Pt(base_font_size)
            r.font.name = FONTS['body']
            r.font.color.rgb = BLACK

        elif match.group('code'):
            r = paragraph.add_run(match.group('code_text'))
            r.font.name = FONTS['code']
            r.font.size = Pt(SIZES['code_pt'])
            r.font.color.rgb = BLACK
            set_run_shading(r, COLORS['inline_code_bg'])

        elif match.group('link'):
            add_hyperlink(paragraph, match.group('link_text'), match.group('link_url'), PRIMARY_BLUE)

        elif match.group('image'):
            alt = match.group('alt')
            img_url = match.group('img_url')
            render_inline_image(paragraph._parent, alt, img_url)

        last_end = match.end()

    # Add remaining plain text after last match
    if last_end < len(text):
        tail = text[last_end:]
        r = paragraph.add_run(tail)
        r.font.size = Pt(base_font_size)
        r.font.name = FONTS['body']
        r.font.color.rgb = BLACK


# ============================================================================
# LAYER 4: BLOCK PARSER (STATE MACHINE)
# ============================================================================

def preprocess_markdown(raw_text):
    """Remove YAML frontmatter from markdown text."""
    return re.sub(r'^---\s*\n.*?\n---\s*\n', '', raw_text, flags=re.DOTALL)


def parse_col_alignment(sep_cells):
    """Parse table separator row to determine column alignments."""
    alignments = []
    for cell in sep_cells:
        cell = cell.strip()
        if not cell or cell == '---':
            alignments.append('left')
        elif cell.startswith(':') and cell.endswith(':'):
            alignments.append('center')
        elif cell.endswith(':'):
            alignments.append('right')
        else:
            alignments.append('left')
    return alignments


def is_numeric_column(col_values):
    """Detect if a column appears to be numeric."""
    if not col_values:
        return False

    numeric_count = 0
    for val in col_values:
        val = val.strip()
        if not val:
            continue
        # Strip currency and percentage symbols
        test_val = val.replace('%', '').replace('$', '').replace(',', '').replace('AUD', '').strip()
        try:
            float(test_val)
            numeric_count += 1
        except ValueError:
            pass

    return numeric_count > 0 and numeric_count / max(len(col_values), 1) > 0.7


def classify_line(line):
    """Classify a single line of markdown and return (type, *args) tuple."""
    stripped = line.strip()

    # Empty line
    if not stripped:
        return ('blank',)

    # Heading
    heading_match = re.match(r'^(#{1,3})\s+(.+)$', line)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2).strip()
        return ('heading', level, text)

    # Callout marker
    callout_match = re.match(r'^>\s*\[!([A-Z]+)\]', line)
    if callout_match:
        kind = callout_match.group(1)
        if kind in CALLOUT_COLORS:
            return ('callout_marker', kind)

    # Blockquote (non-callout)
    if stripped.startswith('>'):
        content = stripped[1:].strip()
        return ('blockquote', content)

    # Code fence
    if re.match(r'^```+', stripped):
        lang = re.match(r'^```+(\w*)', stripped)
        lang_name = lang.group(1) if lang and lang.group(1) else 'text'
        return ('code_fence', lang_name)

    # Horizontal rule
    if re.match(r'^(---+|===+|\*\*\*+|___+)$', stripped):
        return ('hr',)

    # Table row
    if '|' in line and stripped.startswith('|'):
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        # Check if this is a separator row
        is_sep = all(re.match(r'^:?-+:?$', cell.strip()) for cell in cells)
        if is_sep:
            return ('table_sep', parse_col_alignment(cells))
        else:
            return ('table_row', cells)

    # Bullet list
    bullet_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
    if bullet_match:
        depth = len(bullet_match.group(1)) // 2 + 1
        text = bullet_match.group(2).strip()
        return ('bullet', depth, text)

    # Numbered list
    numbered_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
    if numbered_match:
        depth = len(numbered_match.group(1)) // 2 + 1
        text = numbered_match.group(2).strip()
        return ('numbered', depth, text)

    # Regular paragraph
    return ('paragraph', stripped)


# ============================================================================
# LAYER 5: BLOCK RENDERERS
# ============================================================================

def render_heading(doc, level, text):
    """Render markdown heading with Word's built-in heading style."""
    p = doc.add_heading(text, level=level)

    # Clear auto-created run and re-add with parse_inline for proper formatting
    p.clear()
    parse_inline(p, text, base_font_size=SIZES.get(f'h{level}_pt', 11))

    # Apply heading style formatting
    for run in p.runs:
        run.font.bold = HEADINGS.get(f'h{level}_bold', True)
        run.font.size = Pt(SIZES[f'h{level}_pt'])

        if level == 1:
            run.font.color.rgb = hex_to_rgbcolor(HEADINGS['h1_color'])
        elif level == 2:
            run.font.color.rgb = hex_to_rgbcolor(HEADINGS['h2_color'])
        else:
            run.font.color.rgb = hex_to_rgbcolor(HEADINGS['h3_color'])

    return p


def render_paragraph(doc, text):
    """Render a regular paragraph with inline formatting."""
    p = doc.add_paragraph()
    parse_inline(p, text, base_font_size=SIZES['body_pt'])

    # Style all runs
    for run in p.runs:
        run.font.size = Pt(SIZES['body_pt'])
        run.font.name = FONTS['body']
        run.font.color.rgb = BLACK

    p.paragraph_format.space_after = Pt(6)
    return p


def render_table(doc, header_row, data_rows, alignments):
    """Render markdown table with styled header and borders."""
    num_cols = len(header_row)
    num_rows = 1 + len(data_rows)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    if TABLE_CFG.get('center_table'):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Ensure alignments list matches column count
    if len(alignments) < num_cols:
        alignments.extend(['left'] * (num_cols - len(alignments)))

    # Auto-detect numeric columns for unspecified alignments
    for col_idx in range(num_cols):
        if col_idx < len(alignments) and alignments[col_idx] == 'left':
            col_values = [row[col_idx] if col_idx < len(row) else '' for row in data_rows]
            if is_numeric_column(col_values):
                alignments[col_idx] = 'center'

    # Style header row
    header_cells = table.rows[0].cells
    for col_idx, cell_text in enumerate(header_row):
        cell = header_cells[col_idx]
        set_cell_background(cell, TABLE_CFG['header_bg'])
        cell.text = ''
        p = cell.paragraphs[0]
        parse_inline(p, cell_text, base_font_size=SIZES['table_header_pt'])
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = hex_to_rgbcolor(TABLE_CFG['header_fg'])
            run.font.size = Pt(SIZES['table_header_pt'])

        # Set alignment
        if col_idx < len(alignments):
            if alignments[col_idx] == 'center':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignments[col_idx] == 'right':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Style data rows
    for row_idx, row_data in enumerate(data_rows, start=1):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            parse_inline(p, cell_text, base_font_size=SIZES['table_body_pt'])
            for run in p.runs:
                run.font.color.rgb = BLACK
                run.font.size = Pt(SIZES['table_body_pt'])

            # Set alignment
            if col_idx < len(alignments):
                if alignments[col_idx] == 'center':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignments[col_idx] == 'right':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Apply borders
    set_table_borders(table)

    return table


def render_callout(doc, kind, lines):
    """Render callout block (NOTE, TIP, WARNING, IMPORTANT)."""
    colors = CALLOUT_COLORS.get(kind, CALLOUT_COLORS['NOTE'])
    border_color = colors['border']
    fill_color = colors['fill']

    for idx, line in enumerate(lines):
        p = doc.add_paragraph()

        # Add label to first line
        if idx == 0:
            label_run = p.add_run(f"{kind}: ")
            label_run.font.bold = True
            label_run.font.size = Pt(SIZES['body_pt'])
            label_run.font.color.rgb = hex_to_rgbcolor(border_color)

        # Add content with inline formatting
        parse_inline(p, line, base_font_size=SIZES['body_pt'])

        # Apply callout styling
        set_paragraph_left_border(p, border_color, sz=24)
        set_paragraph_shading(p, fill_color)
        set_paragraph_indent(p, left_twips=360)

        for run in p.runs:
            run.font.size = Pt(SIZES['body_pt'])
            run.font.name = FONTS['body']
            if run.text != f"{kind}: ":
                run.font.color.rgb = BLACK


def render_blockquote(doc, lines):
    """Render blockquote with grey left border and italic styling."""
    for line in lines:
        p = doc.add_paragraph()

        # Add content with inline formatting
        parse_inline(p, line, base_font_size=SIZES['body_pt'])

        # Apply blockquote styling
        set_paragraph_left_border(p, COLORS['blockquote_border'], sz=12)
        set_paragraph_indent(p, left_twips=360)

        for run in p.runs:
            run.font.italic = True
            run.font.color.rgb = hex_to_rgbcolor(COLORS['blockquote_text'])
            run.font.size = Pt(SIZES['body_pt'])


def render_code_block(doc, lines, lang='text'):
    """Render code block with monospace font and grey background."""
    p = doc.add_paragraph()

    for idx, line in enumerate(lines):
        if idx > 0:
            # Insert line break within the paragraph
            p.add_run('\n')
        run = p.add_run(line)
        run.font.name = FONTS['code']
        run.font.size = Pt(SIZES['code_pt'])
        run.font.color.rgb = BLACK

    # Apply code block styling
    set_paragraph_shading(p, COLORS['code_block_bg'])
    set_paragraph_indent(p, left_twips=360, right_twips=360)

    # Disable spell check
    for run in p.runs:
        run.font.no_proof = True


def render_bullet(doc, text, depth=1):
    """Render bullet list item."""
    style = 'List Bullet' if depth == 1 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    p.text = ''  # Clear auto-added text

    parse_inline(p, text, base_font_size=SIZES['body_pt'])

    for run in p.runs:
        run.font.size = Pt(SIZES['body_pt'])
        run.font.name = FONTS['body']
        run.font.color.rgb = BLACK

    p.paragraph_format.left_indent = Inches(0.25 * depth)
    return p


def render_numbered(doc, text, depth=1):
    """Render numbered list item."""
    style = 'List Number' if depth == 1 else 'List Number 2'
    p = doc.add_paragraph(style=style)
    p.text = ''  # Clear auto-added text

    parse_inline(p, text, base_font_size=SIZES['body_pt'])

    for run in p.runs:
        run.font.size = Pt(SIZES['body_pt'])
        run.font.name = FONTS['body']
        run.font.color.rgb = BLACK

    p.paragraph_format.left_indent = Inches(0.25 * depth)
    return p


def render_inline_image(doc, alt, img_url):
    """Render inline image or alt text fallback."""
    img_path = Path(img_url)

    if img_path.exists():
        # Embed the image
        try:
            doc.add_picture(str(img_path), width=Inches(5))
        except Exception:
            # If image fails to load, use alt text
            p = doc.add_paragraph(f"[Image: {alt}]")
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = hex_to_rgbcolor(COLORS['blockquote_text'])
                run.font.size = Pt(SIZES['caption_pt'])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Image doesn't exist, show alt text
        p = doc.add_paragraph(f"[Image: {alt}]")
        for run in p.runs:
            run.font.italic = True
            run.font.color.rgb = hex_to_rgbcolor(COLORS['blockquote_text'])
            run.font.size = Pt(SIZES['caption_pt'])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ============================================================================
# LAYER 6: MAIN ORCHESTRATION
# ============================================================================

def convert_md_to_docx(input_path, output_path=None, config_path=None):
    """Convert markdown file to styled Word document."""

    # Resolve input path
    input_path = Path(input_path)
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    # Determine output path
    if output_path is None:
        output_path = input_path.parent / f"{input_path.stem}_branded.docx"
    else:
        output_path = Path(output_path)

    # Read markdown content
    with open(input_path, 'r', encoding='utf-8') as f:
        raw_text = f.read()

    # Preprocess
    md_text = preprocess_markdown(raw_text)
    lines = md_text.split('\n')

    # Create document
    doc = Document()

    # Set page layout
    section = doc.sections[0]
    margin = Inches(PAGE_CFG['margin_inches'])
    section.left_margin = margin
    section.right_margin = margin
    section.top_margin = margin
    section.bottom_margin = margin
    section.page_width = Inches(PAGE_CFG['width_inches'])
    section.page_height = Inches(PAGE_CFG['height_inches'])

    # Set default font
    style = doc.styles['Normal']
    style.font.name = FONTS['body']
    style.font.size = Pt(SIZES['body_pt'])
    style.font.color.rgb = BLACK

    # State machine
    in_table = False
    in_code_block = False
    in_callout = False
    in_blockquote = False

    table_header = []
    table_rows = []
    table_alignments = []
    code_lines = []
    code_lang = 'text'
    callout_lines = []
    callout_kind = ''
    blockquote_lines = []

    def flush_table():
        nonlocal in_table, table_header, table_rows, table_alignments
        if in_table and table_header:
            render_table(doc, table_header, table_rows, table_alignments)
            doc.add_paragraph()  # Spacing
        in_table = False
        table_header = []
        table_rows = []
        table_alignments = []

    def flush_code_block():
        nonlocal in_code_block, code_lines, code_lang
        if in_code_block and code_lines:
            render_code_block(doc, code_lines, code_lang)
            doc.add_paragraph()  # Spacing
        in_code_block = False
        code_lines = []
        code_lang = 'text'

    def flush_callout():
        nonlocal in_callout, callout_lines, callout_kind
        if in_callout and callout_lines:
            render_callout(doc, callout_kind, callout_lines)
            doc.add_paragraph()  # Spacing
        in_callout = False
        callout_lines = []
        callout_kind = ''

    def flush_blockquote():
        nonlocal in_blockquote, blockquote_lines
        if in_blockquote and blockquote_lines:
            render_blockquote(doc, blockquote_lines)
            doc.add_paragraph()  # Spacing
        in_blockquote = False
        blockquote_lines = []

    # Process lines
    i = 0
    while i < len(lines):
        line = lines[i]
        token = classify_line(line)

        token_type = token[0]

        if token_type == 'heading':
            flush_table()
            flush_code_block()
            flush_callout()
            flush_blockquote()
            _, level, text = token
            render_heading(doc, level, text)

        elif token_type == 'table_row':
            if not in_table:
                flush_code_block()
                flush_callout()
                flush_blockquote()
                in_table = True
            _, cells = token
            if not table_header:
                table_header = cells
            else:
                table_rows.append(cells)

        elif token_type == 'table_sep':
            if in_table:
                _, alignments = token
                table_alignments = alignments

        elif token_type == 'code_fence':
            if in_code_block:
                flush_code_block()
            else:
                flush_table()
                flush_callout()
                flush_blockquote()
                in_code_block = True
                _, lang = token
                code_lang = lang
            i += 1
            continue

        elif token_type == 'callout_marker':
            flush_table()
            flush_code_block()
            flush_blockquote()
            _, kind = token
            in_callout = True
            callout_kind = kind

        elif token_type == 'blockquote':
            if in_callout:
                _, content = token
                callout_lines.append(content)
            elif in_blockquote:
                _, content = token
                blockquote_lines.append(content)
            else:
                flush_table()
                flush_code_block()
                in_blockquote = True
                _, content = token
                blockquote_lines.append(content)

        elif token_type == 'hr':
            flush_table()
            flush_code_block()
            flush_callout()
            flush_blockquote()
            add_horizontal_rule(doc)

        elif token_type == 'bullet':
            flush_table()
            flush_code_block()
            flush_callout()
            flush_blockquote()
            _, depth, text = token
            render_bullet(doc, text, depth)

        elif token_type == 'numbered':
            flush_table()
            flush_code_block()
            flush_callout()
            flush_blockquote()
            _, depth, text = token
            render_numbered(doc, text, depth)

        elif token_type == 'blank':
            if in_code_block:
                code_lines.append('')
            # Skip blanks otherwise (natural spacing from paragraph formatting)

        elif token_type == 'paragraph':
            flush_table()
            flush_code_block()
            flush_callout()
            flush_blockquote()
            _, text = token
            if text.strip():
                render_paragraph(doc, text)

        i += 1

    # Flush any remaining state
    flush_table()
    flush_code_block()
    flush_callout()
    flush_blockquote()

    # Save document
    doc.save(str(output_path))
    return str(output_path)


def main():
    """CLI entry point."""
    if len(sys.argv) < 2:
        print("Usage: python generate_branded_report.py <input.md> [output.docx]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    try:
        output_file = convert_md_to_docx(input_path, output_path)
        print(f"Document created: {output_file}")
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
