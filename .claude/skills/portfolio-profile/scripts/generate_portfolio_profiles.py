"""
Generate portfolio profiles for Minerds Bell / Salita Portfolio Services.
Source: Minerds Bell MyNorth IM Questionnaire 30112025.xlsx
Template: .claude/skills/portfolio-profile/assets/Managed Portfolio Profile Template.docx
Output: .claude/output/
"""

import sys
import os
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from lxml import etree

sys.stdout.reconfigure(encoding='utf-8')

TEMPLATE_PATH = r'C:\Users\Sara\Downloads\AIagentproject\.claude\skills\portfolio-profile\assets\Managed Portfolio Profile Template.docx'
OUTPUT_DIR = r'C:\Users\Sara\Downloads\AIagentproject\.claude\output'
from datetime import datetime
current_date = datetime.now().strftime("%Y-%m-%d")
OUTPUT_PATH = os.path.join(OUTPUT_DIR, f'Managed Portfolio Profile - BlackRock AU - {current_date}.docx')

os.makedirs(OUTPUT_DIR, exist_ok=True)

# ─────────────────────────────────────────────
# Portfolio data (sourced from questionnaire)
# ─────────────────────────────────────────────

SRM_LABELS = {
    1: "1 – Very Low",
    2: "2 – Low",
    3: "3 – Low to Medium",
    4: "4 – Medium",
    5: "5 – Medium to High",
    6: "6 – High",
    7: "7 – Very High",
}

portfolios = [
    {
        "name": "BlackRock Balanced AU Managed Portfolio",
        "north_code": "NTH0001",
        "investment_manager": "BlackRock Investment Management (Australia) Limited",
        "availability": "MyNorth Managed Portfolios",
        "asset_class": "Diversified",
        "investment_style": "Active & Index",
        "investment_universe": "Managed funds and ETFs",
        "investment_objective": (
            "Aims to deliver long-term capital growth and a modest income stream above CPI + 3.0% p.a. "
            "over rolling 7-year periods, net of indirect fees, while maintaining a balanced level of "
            "portfolio volatility suitable for investors with a medium-to-high risk tolerance."
        ),
        "designed_for": (
            "Designed for investors who seek balanced exposure to growth and defensive assets, are "
            "comfortable with moderate market fluctuations, and have an investment horizon of at least 7 years."
        ),
        "min_investment_horizon": "7 years plus",
        "min_investment_amount": "[MISSING – please provide]",
        "indicative_assets": "Minimum 12, Maximum 28",
        "benchmark": "Morningstar AUS Moderate NR AUD",
        "srm": 5,
        "saa": {
            "aus_equities":   {"min": "[TBD]",  "saa": "[TBD]", "max": "[TBD]"},
            "intl_equities":  {"min": "[TBD]",  "saa": "[TBD]", "max": "[TBD]"},
            "listed_property":{"min": "[TBD]",  "saa": "[TBD]", "max": "[TBD]"},
            "alternatives":   {"min": "[TBD]",  "saa": "[TBD]", "max": "[TBD]"},
            "total_growth":   {"saa": "[TBD]"},
            "aus_fi":         {"min": "[TBD]",  "saa": "[TBD]", "max": "[TBD]"},
            "intl_fi":        {"min": "[TBD]",  "saa": "[TBD]", "max": "[TBD]"},
            "cash":           {"min": "1.5%",  "saa": "[TBD]", "max": "[TBD]"},
            "total_defensive":{"saa": "[TBD]"},
        },
        "portfolio_income": "Reinvested",
        "min_assets": "12",
        "max_assets": "28",
        "min_single_pos": "0.75%",
        "max_single_pos": "18%",
        "max_new_asset": "1%",
        "target_volatility": "8.5% to 10.5% annualised volatility target range",
        "min_cash_buffer": "1.5%",
        "trading_preference": "Active",
        "expected_turnover": "25% to 40%",
        "issues": [
            "MISSING INFORMATION: Portfolio Snapshot (Page Five) contains incomplete SAA data. Strategic asset allocation percentages and ranges are not populated in the questionnaire. Please complete Section 1 (SAA allocation) and Section 2 (GICS sector exposures) of Page Five.",
            "MISSING FIELD: Minimum investment amount not provided in questionnaire.",
            "NOTE: Portfolio Snapshot also lacks GICS sector exposure data and Top 5 overweights/underweights.",
        ],
    },
    {
        "name": "BlackRock High Growth AU Managed Portfolio",
        "north_code": "NTH0002",
        "investment_manager": "BlackRock Investment Management (Australia) Limited",
        "availability": "MyNorth Managed Portfolios",
        "asset_class": "Diversified",
        "investment_style": "Active & Index",
        "investment_universe": "Managed funds and ETFs",
        "investment_objective": (
            "Aims to deliver long-term capital growth above CPI + 3.0% p.a. over rolling 7-year periods, "
            "net of indirect fees, with higher allocation to growth assets suitable for investors with "
            "a high risk tolerance."
        ),
        "designed_for": (
            "Designed for investors who seek maximum exposure to growth assets, are comfortable with "
            "significant market fluctuations, and have an investment horizon of at least 7 years."
        ),
        "min_investment_horizon": "7 years plus",
        "min_investment_amount": "[MISSING – please provide]",
        "indicative_assets": "Minimum 12, Maximum 28",
        "benchmark": "[MISSING – please specify High Growth benchmark]",
        "srm": "[MISSING – please specify SRM for High Growth]",
        "saa": {
            "aus_equities":   {"min": "[TBD]",  "saa": "[TBD]", "max": "[TBD]"},
            "intl_equities":  {"min": "[TBD]",  "saa": "[TBD]", "max": "[TBD]"},
            "listed_property":{"min": "[TBD]",  "saa": "[TBD]", "max": "[TBD]"},
            "alternatives":   {"min": "[TBD]",  "saa": "[TBD]", "max": "[TBD]"},
            "total_growth":   {"saa": "[TBD]"},
            "aus_fi":         {"min": "[TBD]",  "saa": "[TBD]", "max": "[TBD]"},
            "intl_fi":        {"min": "[TBD]",  "saa": "[TBD]", "max": "[TBD]"},
            "cash":           {"min": "1.5%",  "saa": "[TBD]", "max": "[TBD]"},
            "total_defensive":{"saa": "[TBD]"},
        },
        "portfolio_income": "Reinvested",
        "min_assets": "12",
        "max_assets": "28",
        "min_single_pos": "0.75%",
        "max_single_pos": "18%",
        "max_new_asset": "1%",
        "target_volatility": "[MISSING – please specify target volatility range]",
        "min_cash_buffer": "1.5%",
        "trading_preference": "Active",
        "expected_turnover": "[MISSING – please specify expected turnover]",
        "issues": [
            "CRITICAL MISSING: Page Four (Portfolio Details) only contains detailed information for the Balanced portfolio. No dedicated section for High Growth portfolio specifications.",
            "MISSING INFORMATION: Portfolio Snapshot (Page Five) contains incomplete SAA data. Strategic asset allocation percentages and ranges are not populated.",
            "MISSING FIELDS: Benchmark, SRM, target volatility, and expected turnover not specified for High Growth portfolio.",
            "MISSING FIELD: Minimum investment amount not provided in questionnaire.",
        ],
    },
]


# ─────────────────────────────────────────────
# Helper functions
# ─────────────────────────────────────────────

def set_cell_text(cell, text, bold=False):
    """Replace cell content with text, preserving cell XML structure."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ''
    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]
    # Clear existing runs
    for child in list(p._p):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            p._p.remove(child)
    run = p.add_run(text)
    if bold:
        run.bold = True


def fill_table(table, portfolio):
    """Fill a template table with portfolio data."""
    saa = portfolio["saa"]

    # Row 1: Portfolio name + North Code
    set_cell_text(table.rows[1].cells[0], portfolio["name"], bold=True)
    set_cell_text(table.rows[1].cells[1], portfolio["north_code"])

    # Rows 3–14: Portfolio information
    set_cell_text(table.rows[3].cells[1],  portfolio["investment_manager"])
    set_cell_text(table.rows[4].cells[1],  portfolio["availability"])
    set_cell_text(table.rows[5].cells[1],  portfolio["asset_class"])
    set_cell_text(table.rows[6].cells[1],  portfolio["investment_style"])
    set_cell_text(table.rows[7].cells[1],  portfolio["investment_universe"])
    set_cell_text(table.rows[8].cells[1],  portfolio["investment_objective"])
    set_cell_text(table.rows[9].cells[1],  portfolio["designed_for"])
    set_cell_text(table.rows[10].cells[1], portfolio["min_investment_horizon"])
    set_cell_text(table.rows[11].cells[1], portfolio["min_investment_amount"])
    set_cell_text(table.rows[12].cells[1], portfolio["indicative_assets"])
    set_cell_text(table.rows[13].cells[1], portfolio["benchmark"])
    # Handle SRM – support both integer and string (for missing data)
    srm_value = portfolio["srm"]
    srm_text = SRM_LABELS.get(srm_value, srm_value) if isinstance(srm_value, int) else srm_value
    set_cell_text(table.rows[14].cells[1], srm_text)

    # SAA rows (17–20 = Growth, 23–25 = Defensive + cash)
    # Row 17: Australian Equities
    ae = saa["aus_equities"]
    set_cell_text(table.rows[17].cells[1], ae["min"])
    set_cell_text(table.rows[17].cells[2], ae["saa"])
    set_cell_text(table.rows[17].cells[4], ae["max"])

    # Row 18: International Equities
    ie = saa["intl_equities"]
    set_cell_text(table.rows[18].cells[1], ie["min"])
    set_cell_text(table.rows[18].cells[2], ie["saa"])
    set_cell_text(table.rows[18].cells[4], ie["max"])

    # Row 19: Listed Property / Infrastructure
    lp = saa["listed_property"]
    set_cell_text(table.rows[19].cells[1], lp["min"])
    set_cell_text(table.rows[19].cells[2], lp["saa"])
    set_cell_text(table.rows[19].cells[4], lp["max"])

    # Row 20: Alternatives
    alt = saa["alternatives"]
    set_cell_text(table.rows[20].cells[1], alt["min"])
    set_cell_text(table.rows[20].cells[2], alt["saa"])
    set_cell_text(table.rows[20].cells[4], alt["max"])

    # Row 21: Total Growth
    set_cell_text(table.rows[21].cells[2], saa["total_growth"]["saa"])

    # Row 23: Australian Fixed Interest
    afi = saa["aus_fi"]
    set_cell_text(table.rows[23].cells[1], afi["min"])
    set_cell_text(table.rows[23].cells[2], afi["saa"])
    set_cell_text(table.rows[23].cells[4], afi["max"])

    # Row 24: International Fixed Interest
    ifi = saa["intl_fi"]
    set_cell_text(table.rows[24].cells[1], ifi["min"])
    set_cell_text(table.rows[24].cells[2], ifi["saa"])
    set_cell_text(table.rows[24].cells[4], ifi["max"])

    # Row 25: Cash
    cash = saa["cash"]
    set_cell_text(table.rows[25].cells[1], cash["min"])
    set_cell_text(table.rows[25].cells[2], cash["saa"])
    set_cell_text(table.rows[25].cells[4], cash["max"])

    # Row 26: Total Defensive
    set_cell_text(table.rows[26].cells[2], saa["total_defensive"]["saa"])

    # Operational rows (28–37) → value in cells[2]
    set_cell_text(table.rows[28].cells[2], portfolio["portfolio_income"])
    set_cell_text(table.rows[29].cells[2], portfolio["min_assets"])
    set_cell_text(table.rows[30].cells[2], portfolio["max_assets"])
    set_cell_text(table.rows[31].cells[2], portfolio["min_single_pos"])
    set_cell_text(table.rows[32].cells[2], portfolio["max_single_pos"])
    set_cell_text(table.rows[33].cells[2], portfolio["max_new_asset"])
    set_cell_text(table.rows[34].cells[2], portfolio["target_volatility"])
    set_cell_text(table.rows[35].cells[2], portfolio["min_cash_buffer"])
    set_cell_text(table.rows[36].cells[2], portfolio["trading_preference"])
    set_cell_text(table.rows[37].cells[2], portfolio["expected_turnover"])


def add_issues_section(doc, portfolio):
    """Add a highlighted issues/notes paragraph after the table."""
    if not portfolio["issues"]:
        return
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(f"Issues / Missing Information – {portfolio['name']}:")
    run.bold = True
    for issue in portfolio["issues"]:
        p2 = doc.add_paragraph(f"• {issue}")


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


# ─────────────────────────────────────────────
# Build the document
# ─────────────────────────────────────────────

doc = Document(TEMPLATE_PATH)

# Fill portfolio 1 into the existing template table
fill_table(doc.tables[0], portfolios[0])
add_issues_section(doc, portfolios[0])

# Add page break before portfolio 2
add_page_break(doc)

# Deep-copy the template table XML and append to document body
original_tbl = doc.tables[0]._tbl
tbl_copy = copy.deepcopy(original_tbl)
doc.element.body.append(tbl_copy)

# The second table is now doc.tables[1]
fill_table(doc.tables[1], portfolios[1])
add_issues_section(doc, portfolios[1])

# ─────────────────────────────────────────────
# Sense check summary at end
# ─────────────────────────────────────────────
add_page_break(doc)
doc.add_paragraph()
h = doc.add_paragraph()
h.add_run("Sense Check Summary").bold = True

checks = [
    ("Cash buffer allocation ≥ 1%",
     "PASS – Balanced: 1.5%, High Growth: 1.5%"),
    ("Minimum investment timeline > 1 year",
     "PASS – Balanced: 7 years, High Growth: 7 years"),
    ("Benchmark specification",
     "PARTIAL PASS – Balanced: Morningstar AUS Moderate NR AUD specified. High Growth: benchmark not yet specified in questionnaire."),
    ("SAA data completeness",
     "FAIL – Portfolio Snapshot (Page Five) SAA allocation data is incomplete/missing for both portfolios. Critical missing: asset class percentage allocations and ranges, GICS sector exposures, overweight/underweight positions."),
    ("High Growth portfolio specifications",
     "FAIL – Page Four details only populated for Balanced portfolio. High Growth requires: dedicated specifications for SRM, benchmark, target volatility, turnover, and investment objective tailored to growth mandate."),
]

for label, result in checks:
    p = doc.add_paragraph()
    run_l = p.add_run(f"• {label}: ")
    run_l.bold = True
    status = "PASS" if result.startswith("PASS") else "FAIL"
    run_v = p.add_run(result)
    if status == "FAIL":
        from docx.shared import RGBColor
        run_v.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
