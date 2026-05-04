"""
Generate portfolio profiles for Minerds Bell / Salita Portfolio Services.
Source: Minerds Bell MyNorth IM Questionnaire 30112025.xlsx
Template: .codex/skills/portfolio-profile/assets/Managed Portfolio Profile Template.docx
Output: .codex/output/Minerds Bell - Managed Portfolio Profiles.docx
"""

import sys
import os
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from lxml import etree

sys.stdout.reconfigure(encoding='utf-8')

TEMPLATE_PATH = r'C:\Users\Sara\Downloads\AIagentproject\.codex\skills\portfolio-profile\assets\Managed Portfolio Profile Template.docx'
OUTPUT_DIR = r'C:\Users\Sara\Downloads\AIagentproject\.codex\output'
OUTPUT_PATH = os.path.join(OUTPUT_DIR, 'Minerds Bell - Managed Portfolio Profiles.docx')

os.makedirs(OUTPUT_DIR, exist_ok=True)

# ─────────────────────────────────────────────
# Portfolio data (sourced from questionnaire)
# ─────────────────────────────────────────────

SRM_LABELS = {
    1: "1 – Very Low",
    2: "2 – Low",
    3: "3 – Low to Medium",
    4: "4 – Medium",
    5: "5 – Medium to High",
    6: "6 – High",
    7: "7 – Very High",
}

portfolios = [
    {
        "name": "Aurora Growth Managed Portfolio",
        "north_code": "TBD",
        "investment_manager": "Salita Portfolio Services Pty Ltd",
        "availability": "MyNorth",
        "asset_class": "Diversified",
        "investment_style": "Active & Index",
        "investment_universe": "Managed funds and ETFs",
        "investment_objective": (
            "Aims to achieve a return in excess of CPI + 3.75% p.a. and outperform "
            "the Benchmark over a rolling 10 year period, net of indirect fees."
        ),
        "designed_for": (
            "Designed for investors with a minimum 10 year timeframe, who are comfortable "
            "with high investment risk, and/or require high returns to meet their long term objectives."
        ),
        "min_investment_horizon": "10 years plus",
        "min_investment_amount": "[MISSING – please provide]",
        "indicative_assets": "Minimum 10, Maximum 30",
        "benchmark": "Morningstar AUS Agg Tgt Alloc NR AUD",
        "srm": 6,
        "saa": {
            "aus_equities":   {"min": "0%",  "saa": "36%", "max": "85%"},
            "intl_equities":  {"min": "0%",  "saa": "45%", "max": "85%"},
            "listed_property":{"min": "0%",  "saa": "17%", "max": "35% (domestic only – see note)"},
            "alternatives":   {"min": "0%",  "saa": "0%",  "max": "20%"},
            "total_growth":   {"saa": "98%"},
            "aus_fi":         {"min": "0%",  "saa": "0%",  "max": "19%"},
            "intl_fi":        {"min": "0%",  "saa": "0%",  "max": "19%"},
            "cash":           {"min": "1%",  "saa": "2%",  "max": "20%"},
            "total_defensive":{"saa": "2%"},
        },
        "portfolio_income": "Reinvested",
        "min_assets": "10",
        "max_assets": "30",
        "min_single_pos": "1%",
        "max_single_pos": "30%",
        "max_new_asset": "1%",
        "target_volatility": "N/A",
        "min_cash_buffer": "1%",
        "trading_preference": "Active",
        "expected_turnover": "N/A (unlisted managed funds & ETFs only)",
        "issues": [
            "SENSE CHECK ISSUE: International Property & Infrastructure (SAA 14%) has no minimum or "
            "maximum range specified in the questionnaire. The Excel worksheet flagged this portfolio "
            "as 'Correctly Completed: No'. Please confirm the allowable ranges for International Property.",
            "MISSING FIELD: North Code (to be assigned by platform).",
            "MISSING FIELD: Minimum investment amount not provided.",
        ],
    },
    {
        "name": "Aurora Defensive Managed Portfolio",
        "north_code": "TBD",
        "investment_manager": "Salita Portfolio Services Pty Ltd",
        "availability": "MyNorth",
        "asset_class": "Diversified",
        "investment_style": "Active & Index",
        "investment_universe": "Managed funds and ETFs",
        "investment_objective": (
            "Aims to achieve a return in excess of CPI + 0.5% p.a. and outperform "
            "the Benchmark over a rolling 2 year period, net of indirect fees."
        ),
        "designed_for": (
            "Designed for investors with a minimum 2 year timeframe, who have a low tolerance "
            "to investment risk, and/or require low returns to meet their objectives."
        ),
        "min_investment_horizon": "2 years plus",
        "min_investment_amount": "[MISSING – please provide]",
        "indicative_assets": "Minimum 10, Maximum 30",
        "benchmark": "Morningstar AUS Con Tgt Alloc NR AUD",
        "srm": 3,
        "saa": {
            "aus_equities":   {"min": "0%",  "saa": "0%",  "max": "20%"},
            "intl_equities":  {"min": "0%",  "saa": "0%",  "max": "20%"},
            "listed_property":{"min": "0%",  "saa": "0%",  "max": "20%"},
            "alternatives":   {"min": "0%",  "saa": "0%",  "max": "20%"},
            "total_growth":   {"saa": "0%"},
            "aus_fi":         {"min": "0%",  "saa": "59%", "max": "98%"},
            "intl_fi":        {"min": "0%",  "saa": "39%", "max": "98%"},
            "cash":           {"min": "1%",  "saa": "2%",  "max": "100%"},
            "total_defensive":{"saa": "100%"},
        },
        "portfolio_income": "Reinvested",
        "min_assets": "10",
        "max_assets": "30",
        "min_single_pos": "1%",
        "max_single_pos": "30%",
        "max_new_asset": "1%",
        "target_volatility": "N/A",
        "min_cash_buffer": "1%",
        "trading_preference": "Active",
        "expected_turnover": "N/A (unlisted managed funds & ETFs only)",
        "issues": [
            "MISSING FIELD: North Code (to be assigned by platform).",
            "MISSING FIELD: Minimum investment amount not provided.",
        ],
    },
]


# ─────────────────────────────────────────────
# Helper functions
# ─────────────────────────────────────────────

def set_cell_text(cell, text, bold=False):
    """Replace cell content with text, preserving cell XML structure."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ''
    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]
    # Clear existing runs
    for child in list(p._p):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            p._p.remove(child)
    run = p.add_run(text)
    if bold:
        run.bold = True


def fill_table(table, portfolio):
    """Fill a template table with portfolio data."""
    saa = portfolio["saa"]

    # Row 1: Portfolio name + North Code
    set_cell_text(table.rows[1].cells[0], portfolio["name"], bold=True)
    set_cell_text(table.rows[1].cells[1], portfolio["north_code"])

    # Rows 3–14: Portfolio information
    set_cell_text(table.rows[3].cells[1],  portfolio["investment_manager"])
    set_cell_text(table.rows[4].cells[1],  portfolio["availability"])
    set_cell_text(table.rows[5].cells[1],  portfolio["asset_class"])
    set_cell_text(table.rows[6].cells[1],  portfolio["investment_style"])
    set_cell_text(table.rows[7].cells[1],  portfolio["investment_universe"])
    set_cell_text(table.rows[8].cells[1],  portfolio["investment_objective"])
    set_cell_text(table.rows[9].cells[1],  portfolio["designed_for"])
    set_cell_text(table.rows[10].cells[1], portfolio["min_investment_horizon"])
    set_cell_text(table.rows[11].cells[1], portfolio["min_investment_amount"])
    set_cell_text(table.rows[12].cells[1], portfolio["indicative_assets"])
    set_cell_text(table.rows[13].cells[1], portfolio["benchmark"])
    set_cell_text(table.rows[14].cells[1], SRM_LABELS[portfolio["srm"]])

    # SAA rows (17–20 = Growth, 23–25 = Defensive + cash)
    # Row 17: Australian Equities
    ae = saa["aus_equities"]
    set_cell_text(table.rows[17].cells[1], ae["min"])
    set_cell_text(table.rows[17].cells[2], ae["saa"])
    set_cell_text(table.rows[17].cells[4], ae["max"])

    # Row 18: International Equities
    ie = saa["intl_equities"]
    set_cell_text(table.rows[18].cells[1], ie["min"])
    set_cell_text(table.rows[18].cells[2], ie["saa"])
    set_cell_text(table.rows[18].cells[4], ie["max"])

    # Row 19: Listed Property / Infrastructure
    lp = saa["listed_property"]
    set_cell_text(table.rows[19].cells[1], lp["min"])
    set_cell_text(table.rows[19].cells[2], lp["saa"])
    set_cell_text(table.rows[19].cells[4], lp["max"])

    # Row 20: Alternatives
    alt = saa["alternatives"]
    set_cell_text(table.rows[20].cells[1], alt["min"])
    set_cell_text(table.rows[20].cells[2], alt["saa"])
    set_cell_text(table.rows[20].cells[4], alt["max"])

    # Row 21: Total Growth
    set_cell_text(table.rows[21].cells[2], saa["total_growth"]["saa"])

    # Row 23: Australian Fixed Interest
    afi = saa["aus_fi"]
    set_cell_text(table.rows[23].cells[1], afi["min"])
    set_cell_text(table.rows[23].cells[2], afi["saa"])
    set_cell_text(table.rows[23].cells[4], afi["max"])

    # Row 24: International Fixed Interest
    ifi = saa["intl_fi"]
    set_cell_text(table.rows[24].cells[1], ifi["min"])
    set_cell_text(table.rows[24].cells[2], ifi["saa"])
    set_cell_text(table.rows[24].cells[4], ifi["max"])

    # Row 25: Cash
    cash = saa["cash"]
    set_cell_text(table.rows[25].cells[1], cash["min"])
    set_cell_text(table.rows[25].cells[2], cash["saa"])
    set_cell_text(table.rows[25].cells[4], cash["max"])

    # Row 26: Total Defensive
    set_cell_text(table.rows[26].cells[2], saa["total_defensive"]["saa"])

    # Operational rows (28–37) → value in cells[2]
    set_cell_text(table.rows[28].cells[2], portfolio["portfolio_income"])
    set_cell_text(table.rows[29].cells[2], portfolio["min_assets"])
    set_cell_text(table.rows[30].cells[2], portfolio["max_assets"])
    set_cell_text(table.rows[31].cells[2], portfolio["min_single_pos"])
    set_cell_text(table.rows[32].cells[2], portfolio["max_single_pos"])
    set_cell_text(table.rows[33].cells[2], portfolio["max_new_asset"])
    set_cell_text(table.rows[34].cells[2], portfolio["target_volatility"])
    set_cell_text(table.rows[35].cells[2], portfolio["min_cash_buffer"])
    set_cell_text(table.rows[36].cells[2], portfolio["trading_preference"])
    set_cell_text(table.rows[37].cells[2], portfolio["expected_turnover"])


def add_issues_section(doc, portfolio):
    """Add a highlighted issues/notes paragraph after the table."""
    if not portfolio["issues"]:
        return
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(f"Issues / Missing Information – {portfolio['name']}:")
    run.bold = True
    for issue in portfolio["issues"]:
        p2 = doc.add_paragraph(f"• {issue}")


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


# ─────────────────────────────────────────────
# Build the document
# ─────────────────────────────────────────────

doc = Document(TEMPLATE_PATH)

# Fill portfolio 1 into the existing template table
fill_table(doc.tables[0], portfolios[0])
add_issues_section(doc, portfolios[0])

# Add page break before portfolio 2
add_page_break(doc)

# Deep-copy the template table XML and append to document body
original_tbl = doc.tables[0]._tbl
tbl_copy = copy.deepcopy(original_tbl)
doc.element.body.append(tbl_copy)

# The second table is now doc.tables[1]
fill_table(doc.tables[1], portfolios[1])
add_issues_section(doc, portfolios[1])

# ─────────────────────────────────────────────
# Sense check summary at end
# ─────────────────────────────────────────────
add_page_break(doc)
doc.add_paragraph()
h = doc.add_paragraph()
h.add_run("Sense Check Summary").bold = True

checks = [
    ("Cash allocation > 1%",
     "PASS – Aurora Growth: 2%, Aurora Defensive: 2%"),
    ("Asset allocation sums to 100%",
     "PASS – Aurora Growth: 36+45+3+14+0+0+0+2 = 100%; Aurora Defensive: 0+0+0+0+59+39+2 = 100%"),
    ("Minimum investment timeline > 1 year",
     "PASS – Aurora Growth: 10 years, Aurora Defensive: 2 years"),
    ("Benchmark is single-type (Morningstar / RBA+x% / Bloomberg / CPI+x%)",
     "PASS – Both benchmarks are single Morningstar indices"),
    ("Aurora Growth – International Property ranges",
     "FAIL – International Property & Infrastructure (14% SAA) has no min/max ranges specified. "
     "Spreadsheet flagged 'Correctly Completed: No'. Action required from investment manager."),
]

for label, result in checks:
    p = doc.add_paragraph()
    run_l = p.add_run(f"• {label}: ")
    run_l.bold = True
    status = "PASS" if result.startswith("PASS") else "FAIL"
    run_v = p.add_run(result)
    if status == "FAIL":
        from docx.shared import RGBColor
        run_v.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

doc.save(OUTPUT_PATH)
print(f"Document saved to: {OUTPUT_PATH}")
